from __future__ import annotations

import argparse
import json
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING

import pandas as pd

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument


REPO_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_DESCSTAT_CONFIG = REPO_ROOT / "config" / "descstat_variables.csv"
DEFAULT_PANEL = (
    REPO_ROOT
    / "outputs"
    / "analysis_panel"
    / "public_private_nonprofit"
    / "analysis_panel_coa_headroom_2009_2023_public_private_nonprofit.parquet"
)


@dataclass(frozen=True)
class DescstatSpec:
    varname: str
    label: str
    section: str
    units: str
    winsorize: bool
    winsor_lower: float | None
    winsor_upper: float | None
    include_paper: bool
    include_appendix: bool


def parse_bool(value: object) -> bool:
    text = str(value).strip().lower()
    return text in {"1", "true", "yes", "y"}


def parse_optional_float(value: object) -> float | None:
    if pd.isna(value) or str(value).strip() == "":
        return None
    return float(value)


def load_descstat_specs(path: Path = DEFAULT_DESCSTAT_CONFIG) -> list[DescstatSpec]:
    if not path.exists():
        raise FileNotFoundError(f"Descriptive-statistics config not found: {path}")
    df = pd.read_csv(path)
    required = {
        "varname",
        "label",
        "section",
        "units",
        "winsorize",
        "winsor_lower",
        "winsor_upper",
        "include_paper",
        "include_appendix",
    }
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Descriptive-statistics config is missing columns: {', '.join(sorted(missing))}")

    specs: list[DescstatSpec] = []
    seen: set[str] = set()
    for row in df.to_dict("records"):
        varname = str(row["varname"]).strip()
        if not varname:
            continue
        key = varname.upper()
        if key in seen:
            raise ValueError(f"Duplicate descriptive-statistics variable: {varname}")
        seen.add(key)
        lower = parse_optional_float(row["winsor_lower"])
        upper = parse_optional_float(row["winsor_upper"])
        winsorize = parse_bool(row["winsorize"])
        if winsorize and (lower is None or upper is None or not 0 <= lower < upper <= 1):
            raise ValueError(f"Invalid winsorization bounds for {varname}")
        specs.append(
            DescstatSpec(
                varname=varname,
                label=str(row["label"]).strip(),
                section=str(row["section"]).strip(),
                units=str(row["units"]).strip(),
                winsorize=winsorize,
                winsor_lower=lower,
                winsor_upper=upper,
                include_paper=parse_bool(row["include_paper"]),
                include_appendix=parse_bool(row["include_appendix"]),
            )
        )
    return specs


def winsorize_series(series: pd.Series, lower: float | None, upper: float | None) -> tuple[pd.Series, float | None, float | None, int, int]:
    # Winsorization is only for table display. The source panel is never overwritten.
    numeric = pd.to_numeric(series, errors="coerce")
    nonnull = numeric.dropna()
    if nonnull.empty or lower is None or upper is None:
        return numeric, None, None, 0, 0
    lower_cap = float(nonnull.quantile(lower))
    upper_cap = float(nonnull.quantile(upper))
    winsorized = numeric.clip(lower=lower_cap, upper=upper_cap)
    return (
        winsorized,
        lower_cap,
        upper_cap,
        int(numeric.lt(lower_cap).sum()),
        int(numeric.gt(upper_cap).sum()),
    )


def summarize_variable(df: pd.DataFrame, spec: DescstatSpec, order: int) -> dict[str, object]:
    # Each row in the output describes one configured variable from descstat_variables.csv.
    if spec.varname not in df.columns:
        return {
            "order": order,
            "section": spec.section,
            "varname": spec.varname,
            "label": spec.label,
            "units": spec.units,
            "present": False,
        }

    raw = pd.to_numeric(df[spec.varname], errors="coerce")
    winsorized, lower_cap, upper_cap, capped_low, capped_high = winsorize_series(
        raw,
        spec.winsor_lower if spec.winsorize else None,
        spec.winsor_upper if spec.winsorize else None,
    )
    nonnull = raw.dropna()
    quantiles = nonnull.quantile([0.01, 0.25, 0.5, 0.75, 0.99]) if not nonnull.empty else pd.Series(dtype="Float64")

    def value_or_none(series: pd.Series, method: str) -> float | None:
        clean = series.dropna()
        if clean.empty:
            return None
        return float(getattr(clean, method)())

    return {
        "order": order,
        "section": spec.section,
        "varname": spec.varname,
        "label": spec.label,
        "units": spec.units,
        "present": True,
        "winsorized": spec.winsorize,
        "winsor_lower_quantile": spec.winsor_lower,
        "winsor_upper_quantile": spec.winsor_upper,
        "winsor_lower_cap": lower_cap,
        "winsor_upper_cap": upper_cap,
        "n": int(nonnull.count()),
        "missing": int(raw.isna().sum()),
        "missing_share": float(raw.isna().mean()) if len(raw) else 0.0,
        "raw_mean": value_or_none(raw, "mean"),
        "raw_sd": value_or_none(raw, "std"),
        "raw_min": value_or_none(raw, "min"),
        "raw_p01": None if quantiles.empty else float(quantiles.loc[0.01]),
        "raw_p25": None if quantiles.empty else float(quantiles.loc[0.25]),
        "raw_p50": None if quantiles.empty else float(quantiles.loc[0.5]),
        "raw_p75": None if quantiles.empty else float(quantiles.loc[0.75]),
        "raw_p99": None if quantiles.empty else float(quantiles.loc[0.99]),
        "raw_max": value_or_none(raw, "max"),
        "winsor_mean": value_or_none(winsorized, "mean"),
        "winsor_sd": value_or_none(winsorized, "std"),
        "winsor_min": value_or_none(winsorized, "min"),
        "winsor_p50": value_or_none(winsorized, "median"),
        "winsor_max": value_or_none(winsorized, "max"),
        "capped_low": capped_low,
        "capped_high": capped_high,
        "capped_total": capped_low + capped_high,
        "capped_share": float((capped_low + capped_high) / nonnull.count()) if nonnull.count() else 0.0,
    }


def build_descstat_frame(df: pd.DataFrame, specs: list[DescstatSpec]) -> pd.DataFrame:
    rows = [summarize_variable(df, spec, order) for order, spec in enumerate(specs, start=1)]
    return pd.DataFrame(rows).sort_values("order").reset_index(drop=True)


def paper_table(desc: pd.DataFrame) -> pd.DataFrame:
    table = desc[desc["present"].eq(True)].copy()
    table["display_mean"] = table["raw_mean"].where(~table["winsorized"].astype(bool), table["winsor_mean"])
    table["display_sd"] = table["raw_sd"].where(~table["winsorized"].astype(bool), table["winsor_sd"])
    table["tail_handling"] = table.apply(
        lambda row: (
            f"Winsorized p{int(float(row['winsor_lower_quantile']) * 100)}-p{int(float(row['winsor_upper_quantile']) * 100)}"
            if bool(row["winsorized"])
            else "Raw"
        ),
        axis=1,
    )
    table["missing_percent"] = table["missing_share"] * 100
    cols = [
        "section",
        "label",
        "n",
        "missing_percent",
        "display_mean",
        "display_sd",
        "raw_p50",
        "raw_p25",
        "raw_p75",
        "tail_handling",
    ]
    return table[cols].rename(
        columns={
            "section": "Section",
            "label": "Variable",
            "n": "N",
            "missing_percent": "Missing %",
            "display_mean": "Mean",
            "display_sd": "SD",
            "raw_p50": "Median",
            "raw_p25": "p25",
            "raw_p75": "p75",
            "tail_handling": "Tail handling",
        }
    )


def section_slug(section: str) -> str:
    cleaned = "".join(char.lower() if char.isalnum() else "_" for char in section)
    return "_".join(part for part in cleaned.split("_") if part)


def split_paper_tables(paper: pd.DataFrame) -> dict[str, pd.DataFrame]:
    # Smaller section tables are easier to paste into a paper than one wide table.
    return {
        section_slug(section): rows.reset_index(drop=True)
        for section, rows in paper.groupby("Section", sort=False)
    }


def winsor_audit_table(desc: pd.DataFrame) -> pd.DataFrame:
    # This table preserves the raw-versus-capped comparison without cluttering the manuscript table.
    table = desc[desc["present"].eq(True) & desc["winsorized"].astype(bool)].copy()
    table["capped_percent"] = table["capped_share"] * 100
    cols = [
        "section",
        "label",
        "units",
        "raw_mean",
        "winsor_mean",
        "raw_sd",
        "winsor_sd",
        "winsor_lower_cap",
        "winsor_upper_cap",
        "capped_total",
        "capped_percent",
    ]
    return table[cols].rename(
        columns={
            "section": "Section",
            "label": "Variable",
            "units": "Units",
            "raw_mean": "Raw mean",
            "winsor_mean": "Winsorized mean",
            "raw_sd": "Raw SD",
            "winsor_sd": "Winsorized SD",
            "winsor_lower_cap": "Lower cap",
            "winsor_upper_cap": "Upper cap",
            "capped_total": "Rows capped",
            "capped_percent": "Rows capped %",
        }
    )


def appendix_table(desc: pd.DataFrame) -> pd.DataFrame:
    # The appendix table keeps the full audit detail for missingness and tails.
    cols = [
        "section",
        "label",
        "units",
        "n",
        "missing_share",
        "raw_min",
        "raw_p01",
        "raw_p25",
        "raw_p50",
        "raw_p75",
        "raw_p99",
        "raw_max",
        "winsor_lower_cap",
        "winsor_upper_cap",
        "winsor_mean",
        "winsor_sd",
        "capped_low",
        "capped_high",
    ]
    table = desc[desc["present"].eq(True)].copy()
    return table[cols].rename(
        columns={
            "section": "Section",
            "label": "Variable",
            "units": "Units",
            "n": "N",
            "missing_share": "Missing share",
            "raw_min": "Min",
            "raw_p01": "p1",
            "raw_p25": "p25",
            "raw_p50": "Median",
            "raw_p75": "p75",
            "raw_p99": "p99",
            "raw_max": "Max",
            "winsor_lower_cap": "Lower cap",
            "winsor_upper_cap": "Upper cap",
            "winsor_mean": "Mean, winsorized",
            "winsor_sd": "SD, winsorized",
            "capped_low": "Rows capped low",
            "capped_high": "Rows capped high",
        }
    )


LATEX_ESCAPE = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
    "<": r"$<$",
    ">": r"$>$",
}


def latex_escape(value: object) -> str:
    text = "" if pd.isna(value) else str(value)
    return "".join(LATEX_ESCAPE.get(char, char) for char in text)


LEFT_ALIGNED_COLUMNS = {
    "Section",
    "Variable",
    "Units",
    "Model",
    "Role",
    "Term",
    "Outcome",
    "Specification",
    "Measure",
    "Fixed effects",
    "Tail handling",
}


def format_value(value: object, column: str) -> str:
    # Format numbers once here so Word, Markdown, and LaTeX show the same values.
    if pd.isna(value):
        return ""
    if column in {"N", "Clusters", "Rows capped", "Rows capped low", "Rows capped high"}:
        try:
            return f"{int(float(str(value).replace(',', ''))):,}"
        except ValueError:
            return str(value)
    if column in {"p", "Within R2", "Missing share", "Missing %", "Rows capped %"}:
        try:
            return f"{float(value):.3f}"
        except ValueError:
            return str(value)
    if isinstance(value, float):
        return f"{value:,.3f}"
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def format_export_table(table: pd.DataFrame) -> pd.DataFrame:
    formatted = table.copy()
    for col in formatted.columns:
        formatted[col] = formatted[col].map(lambda value, column=col: format_value(value, column))
    return formatted


def latex_column_spec(columns: list[str]) -> str:
    # Text columns stay left-aligned; numeric columns are right-aligned for readability.
    alignments = ["l" if col in LEFT_ALIGNED_COLUMNS else "r" for col in columns]
    return "@{}" + "".join(alignments) + "@{}"


def write_latex_table(path: Path, table: pd.DataFrame, caption: str, label: str, note: str | None = None) -> None:
    # The LaTeX export uses booktabs/longtable so tables can span pages cleanly.
    formatted = format_export_table(table)
    column_spec = latex_column_spec(list(formatted.columns))
    header = " & ".join(latex_escape(col) for col in formatted.columns) + r" \\"
    column_count = len(formatted.columns)
    lines = [
        f"\\begin{{longtable}}{{{column_spec}}}",
        f"\\caption{{{latex_escape(caption)}}}\\label{{{label}}}\\\\",
        r"\toprule",
        header,
        r"\midrule",
        r"\endfirsthead",
        f"\\caption[]{{{latex_escape(caption)} (continued)}}\\\\",
        r"\toprule",
        header,
        r"\midrule",
        r"\endhead",
    ]
    if note:
        lines.extend(
            [
                r"\midrule",
                f"\\multicolumn{{{column_count}}}{{p{{0.98\\linewidth}}}}{{\\footnotesize \\textit{{Notes:}} {latex_escape(note)}}}\\\\",
                r"\endfoot",
            ]
        )
    for row in formatted.itertuples(index=False, name=None):
        lines.append(" & ".join(latex_escape(value) for value in row) + r" \\")
    lines.extend([r"\bottomrule", r"\end{longtable}", ""])
    text = "\n".join(lines)
    path.write_text(text, encoding="utf-8")


def markdown_escape(value: object) -> str:
    return str(value).replace("|", "\\|")


def write_markdown_table(path: Path, table: pd.DataFrame, title: str, note: str) -> None:
    # Markdown is the easiest preview to paste into a draft or share in review.
    formatted = format_export_table(table)
    columns = list(formatted.columns)
    header = "| " + " | ".join(markdown_escape(col) for col in columns) + " |"
    separator = "| " + " | ".join("---" if col in LEFT_ALIGNED_COLUMNS else "---:" for col in columns) + " |"
    rows = [
        "| " + " | ".join(markdown_escape(value) for value in row) + " |"
        for row in formatted.itertuples(index=False, name=None)
    ]
    lines = [f"### {title}", "", header, separator, *rows, "", f"Notes: {note}", ""]
    path.write_text("\n".join(lines), encoding="utf-8")


def write_word_table(path: Path, table: pd.DataFrame, title: str, note: str) -> None:
    # Word output is meant for direct copy/paste into Word or Google Docs.
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("Word export requires python-docx. Install dependencies with `pip install -r requirements.txt`.") from exc

    document: DocxDocument = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(11)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formatted = format_export_table(table)
    word_table = document.add_table(rows=1, cols=len(formatted.columns))
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    word_table.style = "Table Grid"

    header_cells = word_table.rows[0].cells
    for idx, col in enumerate(formatted.columns):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(str(col))
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in formatted.itertuples(index=False, name=None):
        cells = word_table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            run = paragraph.add_run(str(value))
            run.font.name = "Times New Roman"
            run.font.size = Pt(8)
            col_name = formatted.columns[idx]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_name in LEFT_ALIGNED_COLUMNS else WD_ALIGN_PARAGRAPH.RIGHT
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    note_paragraph = document.add_paragraph()
    note_run = note_paragraph.add_run(f"Notes: {note}")
    note_run.font.name = "Times New Roman"
    note_run.font.size = Pt(8)
    document.save(path)


def build_descstat_tables(
    input_panel: Path = DEFAULT_PANEL,
    output_dir: Path = Path("outputs/descriptive_tables"),
    config: Path = DEFAULT_DESCSTAT_CONFIG,
    scope_label: str = "public_private_nonprofit",
) -> dict[str, Path]:
    # This is the single public function used by the script and the notebook.
    if not input_panel.exists():
        raise SystemExit(f"Input panel does not exist: {input_panel}")
    specs = load_descstat_specs(config)
    df = pd.read_parquet(input_panel)
    desc = build_descstat_frame(df, specs)
    paper = paper_table(desc[desc["varname"].isin([spec.varname for spec in specs if spec.include_paper])])
    appendix = appendix_table(desc[desc["varname"].isin([spec.varname for spec in specs if spec.include_appendix])])

    scoped_output = output_dir / scope_label
    scoped_output.mkdir(parents=True, exist_ok=True)
    paths = {
        "full_descstat": scoped_output / "descstat_full_pre_post_winsor.csv",
        "paper_csv": scoped_output / "descstat_manuscript_overview.csv",
        "paper_md": scoped_output / "descstat_manuscript_overview.md",
        "paper_tex": scoped_output / "descstat_manuscript_overview.tex",
        "paper_docx": scoped_output / "descstat_manuscript_overview.docx",
        "winsor_audit_csv": scoped_output / "descstat_winsorization_audit.csv",
        "winsor_audit_md": scoped_output / "descstat_winsorization_audit.md",
        "winsor_audit_tex": scoped_output / "descstat_winsorization_audit.tex",
        "winsor_audit_docx": scoped_output / "descstat_winsorization_audit.docx",
        "appendix_csv": scoped_output / "descstat_appendix_distribution_audit.csv",
        "appendix_md": scoped_output / "descstat_appendix_distribution_audit.md",
        "appendix_tex": scoped_output / "descstat_appendix_distribution_audit.tex",
        "appendix_docx": scoped_output / "descstat_appendix_distribution_audit.docx",
        "summary": scoped_output / "descstat_summary.json",
    }
    section_tables = split_paper_tables(paper)
    for slug in section_tables:
        paths[f"section_{slug}_csv"] = scoped_output / f"descstat_section_{slug}.csv"
        paths[f"section_{slug}_md"] = scoped_output / f"descstat_section_{slug}.md"
        paths[f"section_{slug}_tex"] = scoped_output / f"descstat_section_{slug}.tex"
        paths[f"section_{slug}_docx"] = scoped_output / f"descstat_section_{slug}.docx"
    desc.to_csv(paths["full_descstat"], index=False)
    paper.to_csv(paths["paper_csv"], index=False)
    winsor_audit = winsor_audit_table(desc)
    winsor_audit.to_csv(paths["winsor_audit_csv"], index=False)
    appendix.to_csv(paths["appendix_csv"], index=False)
    manuscript_note = (
        "Means and standard deviations use the display rule in the Tail handling column. "
        "Winsorization is for tables only; the analysis panel is unchanged."
    )
    audit_note = (
        "This audit compares raw and winsorized values for variables with p1-p99 caps in config/descstat_variables.csv. "
        "Model estimation uses the prepared panel, not this display table."
    )
    write_markdown_table(paths["paper_md"], paper, "Descriptive statistics: manuscript overview", manuscript_note)
    write_markdown_table(paths["winsor_audit_md"], winsor_audit, "Winsorization audit for descriptive statistics", audit_note)
    write_markdown_table(paths["appendix_md"], appendix, "Appendix distribution audit for descriptive statistics", audit_note)
    write_latex_table(paths["paper_tex"], paper, "Descriptive statistics: manuscript overview", "tab:descstat_overview", manuscript_note)
    write_latex_table(paths["winsor_audit_tex"], winsor_audit, "Winsorization audit for descriptive statistics", "tab:descstat_winsor_audit", audit_note)
    write_latex_table(paths["appendix_tex"], appendix, "Appendix distribution audit for descriptive statistics", "tab:appendix_descstat_distribution", audit_note)
    write_word_table(paths["paper_docx"], paper, "Descriptive statistics: manuscript overview", manuscript_note)
    write_word_table(paths["winsor_audit_docx"], winsor_audit, "Winsorization audit for descriptive statistics", audit_note)
    write_word_table(paths["appendix_docx"], appendix, "Appendix distribution audit for descriptive statistics", audit_note)
    for slug, table in section_tables.items():
        title = f"Descriptive statistics: {table['Section'].iloc[0]}"
        table.to_csv(paths[f"section_{slug}_csv"], index=False)
        write_markdown_table(paths[f"section_{slug}_md"], table, title, manuscript_note)
        write_latex_table(paths[f"section_{slug}_tex"], table, title, f"tab:descstat_{slug}", manuscript_note)
        write_word_table(paths[f"section_{slug}_docx"], table, title, manuscript_note)
    summary = {
        "built_at_utc": datetime.now(timezone.utc).isoformat(),
        "input_panel": str(input_panel),
        "config": str(config),
        "scope_label": scope_label,
        "rows": int(len(df)),
        "columns": int(len(df.columns)),
        "configured_variables": int(len(specs)),
        "present_variables": int(desc["present"].fillna(False).astype(bool).sum()),
        "paper_rows": int(len(paper)),
        "section_tables": {slug: int(len(table)) for slug, table in section_tables.items()},
        "winsor_audit_rows": int(len(winsor_audit)),
        "appendix_rows": int(len(appendix)),
        "outputs": {key: str(value) for key, value in paths.items() if key != "summary"},
    }
    paths["summary"].write_text(json.dumps(summary, indent=2, sort_keys=True), encoding="utf-8")
    return paths


def main() -> None:
    parser = argparse.ArgumentParser(description="Build descriptive-statistics tables before and after winsorization.")
    parser.add_argument("--input-panel", type=Path, default=DEFAULT_PANEL)
    parser.add_argument("--output-dir", type=Path, default=Path("outputs/descriptive_tables"))
    parser.add_argument("--config", type=Path, default=DEFAULT_DESCSTAT_CONFIG)
    parser.add_argument("--scope-label", default="public_private_nonprofit")
    args = parser.parse_args()
    paths = build_descstat_tables(args.input_panel, args.output_dir, args.config, args.scope_label)
    print(f"Wrote {paths['summary'].parent}")


if __name__ == "__main__":
    main()
